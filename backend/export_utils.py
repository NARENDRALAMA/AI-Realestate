"""Generate downloadable files (TXT, PDF, DOCX) from a content bundle."""
from __future__ import annotations

import io

from models import GeneratedBundle

_SECTIONS = [
    ("listing", "Listing Description"),
    ("social", "Social Media Post"),
    ("email", "Email Copy"),
    ("video", "Video Script"),
]


# ─── TXT ──────────────────────────────────────────────────────────────────────

def build_txt(bundle: GeneratedBundle, title: str, tone: str) -> bytes:
    lines: list[str] = [
        "REM CONTENT STUDIO",
        f"Property: {title}",
        f"Tone: {tone.title()}",
        "=" * 60,
    ]
    for key, heading in _SECTIONS:
        text = getattr(bundle, key)
        if text:
            lines += ["", f"## {heading}", "", text, "", "-" * 60]
    return "\n".join(lines).encode("utf-8")


# ─── PDF ──────────────────────────────────────────────────────────────────────

def build_pdf(bundle: GeneratedBundle, title: str, tone: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()

    # Custom style for body text that handles preformatted lines
    body_style = ParagraphStyle(
        "body",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=4,
    )
    heading2_style = ParagraphStyle(
        "h2",
        parent=styles["Heading2"],
        fontSize=13,
        spaceBefore=14,
        spaceAfter=6,
    )

    story = []
    story.append(Paragraph(f"REM Content Studio — {title}", styles["Title"]))
    story.append(Paragraph(f"Tone: {tone.title()}", styles["Italic"]))
    story.append(Spacer(1, 0.6 * cm))

    for key, heading in _SECTIONS:
        text: str = getattr(bundle, key)
        if not text:
            continue
        story.append(Paragraph(heading, heading2_style))
        for line in text.split("\n"):
            safe = (
                line
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            if safe.strip():
                story.append(Paragraph(safe, body_style))
            else:
                story.append(Spacer(1, 0.2 * cm))
        story.append(Spacer(1, 0.5 * cm))

    doc.build(story)
    return buffer.getvalue()


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def build_docx(bundle: GeneratedBundle, title: str, tone: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    t = doc.add_heading(f"REM Content Studio — {title}", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph(f"Tone: {tone.title()}")
    sub.runs[0].italic = True

    doc.add_paragraph()

    for key, heading in _SECTIONS:
        text: str = getattr(bundle, key)
        if not text:
            continue
        doc.add_heading(heading, level=1)
        for line in text.split("\n"):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
